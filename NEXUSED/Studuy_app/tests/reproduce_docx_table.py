import docx
import os
import sys

# Mocking the extraction logic from backend_logic.py
def extract_text_from_docx_current(filepath):
    doc = docx.Document(filepath)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text

def create_test_docx_with_table(filename):
    doc = docx.Document()
    doc.add_paragraph("This is a paragraph before the table.")
    
    table = doc.add_table(rows=2, cols=2)
    cell = table.cell(0, 0)
    cell.text = "Header 1"
    cell = table.cell(0, 1)
    cell.text = "Header 2"
    cell = table.cell(1, 0)
    cell.text = "Row 1, Col 1"
    cell = table.cell(1, 1)
    cell.text = "Row 1, Col 2"
    
    doc.add_paragraph("This is a paragraph after the table.")
    doc.save(filename)
    print(f"Created {filename} with a table.")

def test_extraction():
    filename = "test_table.docx"
    try:
        create_test_docx_with_table(filename)
        print("Extracting text using current logic...")
        extracted_text = extract_text_from_docx_current(filename)
        
        print("-" * 20)
        print("EXTRACTED TEXT:")
        print(extracted_text)
        print("-" * 20)
        
        if "Header 1" in extracted_text:
            print("SUCCESS: Table content extraction successful (Unexpected for current logic).")
        else:
            print("FAILURE: Table content MISSING. Current logic ignores tables.")
            
    except Exception as e:
        print(f"Error: {e}")
    finally:
        if os.path.exists(filename):
            os.remove(filename)

if __name__ == "__main__":
    test_extraction()
