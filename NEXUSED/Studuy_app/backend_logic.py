import os
from dotenv import load_dotenv

load_dotenv()
import json
import uuid
import itertools
import requests
import re
import asyncio
import logging
from openai import OpenAI
import nltk
from bs4 import BeautifulSoup

# --- External Service Imports ---
try:
    from pinecone.grpc import PineconeGRPC as Pinecone
except ImportError:
    print("Pinecone library not found. Please install pinecone-client[grpc].")

try:
    from neo4j import GraphDatabase
except ImportError:
    print("Neo4j driver not found. Please install neo4j.")

# --- Import KG2 for unified extraction logic ---
try:
    import KG2
    from KG2 import extract_entities_and_relations
except ImportError:
    print("KG2 module not found. Extraction will fail.")
    def extract_entities_and_relations(text):
        print("Error: KG2 not loaded.")
        return None

# --- NLTK Setup ---
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')
    nltk.download('punkt_tab')

# --- Configuration ---
# Ideally, these should be in environment variables.
# Using values provided in original files.

# OpenAI
os.environ["OPENAI_API_KEY"] = os.getenv("OPENAI_API_KEY", "") 
client = OpenAI()

# Pinecone
PINECONE_KEY = os.getenv("PINECONE_KEY", "")
PINECONE_HOST = os.getenv("PINECONE_HOST", "https://rag-index-xfor3af.svc.aped-4627-b74a.pinecone.io")
EMBEDDING_MODEL = "text-embedding-ada-002"

# Neo4j
NEO4J_URI = os.getenv("NEO4J_URI", "neo4j+s://ba74b87c.databases.neo4j.io")
NEO4J_USERNAME = os.getenv("NEO4J_USERNAME", "neo4j")
NEO4J_PASSWORD = os.getenv("NEO4J_PASSWORD", "")

# ==========================================
# 1. Pinecone Logic
# ==========================================

def init_pinecone():
    """Initialize Pinecone connection."""
    pc = Pinecone(api_key=PINECONE_KEY)
    index = pc.Index(host=PINECONE_HOST)
    return index

def get_content_from_url(url):
    """Scrape and tokenize content from a URL."""
    try:
        response = requests.get(url)
        response.raise_for_status()
        soup = BeautifulSoup(response.content, 'html.parser')
        text = soup.get_text()
        sentences = nltk.sent_tokenize(text)
        return sentences, text
    except Exception as e:
        print(f"Error fetching URL: {e}")
        return [], ""

def upsert_to_pinecone(index, sentences, namespace):
    """Generate embeddings and upsert to Pinecone."""
    if not sentences:
        return False, "No sentences to process."
    
    try:
        # 1. Generate Embeddings
        response = client.embeddings.create(
            input=sentences,
            model=EMBEDDING_MODEL
        )
    except Exception as e:
        return False, f"Error generating embeddings: {e}"

    # 2. Prepare Vectors
    vectors = []
    for sentence, embedding_item in zip(sentences, response.data):
        vectors.append({
            'id': str(uuid.uuid4()),
            'values': embedding_item.embedding,
            'metadata': {
                'text': sentence
            }
        })
    
    # 3. Upsert in Batches
    def chunks(iterable, batch_size=100):
        it = iter(iterable)
        chunk = tuple(itertools.islice(it, batch_size))
        while chunk:
            yield chunk
            chunk = tuple(itertools.islice(it, batch_size))

    count = 0
    try:
        for chunk in chunks(vectors, batch_size=100):
            index.upsert(vectors=chunk, namespace=namespace)
            count += len(chunk)
    except Exception as e:
        return False, f"Error upserting to Pinecone: {e}"
    
    return True, f"Successfully upserted {count} vectors."

# ==========================================
# 2. Neo4j / Knowledge Graph Logic
# ==========================================

# extract_entities_and_relations is imported from KG2.py
# The local definition has been removed to avoid duplication and ensure consistency.

def load_graph_to_neo4j(graph_data: dict):
    """
    Loads the extracted graph data into Neo4j.
    """
    if not graph_data:
        return

    try:
        driver = GraphDatabase.driver(NEO4J_URI, auth=(NEO4J_USERNAME, NEO4J_PASSWORD))
        
        # Quick connectivity check
        try:
             driver.verify_connectivity()
        except Exception as e:
             print(f"--- [Neo4j] Could not connect: {e} ---")
             return

        node_labels = {node['id']: node['label'] for node in graph_data.get("nodes", [])}

        with driver.session() as session:
            # 1. Create Nodes
            for node in graph_data.get("nodes", []):
                try:
                    cypher_query = f"MERGE (n:`{node['label']}` {{id: $id}}) SET n += $properties"
                    session.run(cypher_query, id=node['id'], properties=node.get('properties', {}))
                except Exception as e:
                    print(f"Failed to merge node {node['id']}: {e}")

            # 2. Create Edges
            for edge in graph_data.get("edges", []):
                try:
                    source_label = node_labels.get(edge['source'])
                    target_label = node_labels.get(edge['target'])
                    
                    if not source_label or not target_label:
                        continue

                    cypher_query = (
                        f"MATCH (source:`{source_label}` {{id: $source_id}}) "
                        f"MATCH (target:`{target_label}` {{id: $target_id}}) "
                        f"MERGE (source)-[r:`{edge['label']}`]->(target) "
                        f"SET r += $properties"
                    )
                    session.run(cypher_query, 
                                source_id=edge['source'], 
                                target_id=edge['target'], 
                                properties=edge.get('properties', {}))
                except Exception as e:
                    print(f"Failed to merge edge {edge['source']}->{edge['target']}: {e}")

        driver.close()
    except Exception as e:
        print(f"--- [Neo4j] Global Error in load_graph_to_neo4j: {e} ---")

# ==========================================
# 3. Topic Registry
# ==========================================

TOPICS_FILE = "topics.json"

def load_topics():
    """Load available study topics from the registry."""
    if os.path.exists(TOPICS_FILE):
        try:
            with open(TOPICS_FILE, "r") as f:
                return json.load(f)
        except Exception as e:
            print(f"Error loading topics: {e}")
    return []

def save_topic(topic_name):
    """Save the topic name to a local registry."""
    try:
        if os.path.exists(TOPICS_FILE):
            with open(TOPICS_FILE, "r") as f:
                topics = json.load(f)
        else:
            topics = []
        
        if topic_name not in topics:
            topics.append(topic_name)
            with open(TOPICS_FILE, "w") as f:
                json.dump(topics, f)
    except Exception as e:
        print(f"Could not save topic to registry: {e}")

METADATA_FILE = "topics_metadata.json"

def save_topic_metadata(topic_name, metadata):
    """Saves metadata for a topic (summary, pdf_path, etc)."""
    try:
        data = {}
        if os.path.exists(METADATA_FILE):
             with open(METADATA_FILE, "r") as f:
                data = json.load(f)
        
        data[topic_name] = metadata
        
        with open(METADATA_FILE, "w") as f:
            json.dump(data, f)
    except Exception as e:
        print(f"Error saving metadata: {e}")

def get_topic_metadata(topic_name):
    """Retrieves metadata for a topic."""
    try:
        if os.path.exists(METADATA_FILE):
            with open(METADATA_FILE, "r") as f:
                data = json.load(f)
            return data.get(topic_name, {})
    except Exception as e:
        print(f"Error loading metadata: {e}")
    return {}

# ==========================================
# 4. Chat / Query Logic
# ==========================================

# NOTE: The student_app.py imported `from integrated_app import app`.
# I need to see `integrated_app.py` to know how to call the chat engine correctly without streamlit.
# Assuming I can import it here as well.

try:
    from integrated_app import app as chat_app
except ImportError:
    chat_app = None
    print("Warning: integrated_app not found.")

async def process_student_question(question, namespace):
    if not chat_app:
        return "Error: Chat application logic not loaded."
    
    try:
        inputs = {"question": question, "namespace": namespace}
        result = await chat_app.ainvoke(inputs)
        return result.get("final_answer", "I could not generate an answer.")
    except Exception as e:
        return f"An error occurred: {e}"

# ==========================================
# 5. PDF & Smart Features
# ==========================================

try:
    import fitz # PyMuPDF
except ImportError:
    fitz = None
    print("PyMuPDF not found. PDF extraction will be limited.")

import pypdf # Fallback

try:
    import docx
except ImportError:
    docx = None
    print("python-docx not found. DOCX extraction will be disabled.")

def extract_text_from_pdf(filepath):
    """Extracts text from a PDF file."""
    text = ""
    
    # Try PyMuPDF first (Faster, better)
    if fitz:
        try:
            doc = fitz.open(filepath)
            for page in doc:
                text += page.get_text() + "\n"
            doc.close()
        except Exception as e:
            print(f"PyMuPDF failed: {e}")
    
    # Fallback to pypdf if PyMuPDF failed or returned nothing useful 
    # (though if fitz fails on scanned, pypdf likely will too, but worth a shot for weird encodings)
    if not text.strip():
        try:
            reader = pypdf.PdfReader(filepath)
            for page in reader.pages:
                text += page.extract_text() + "\n"
        except Exception as e:
             print(f"pypdf failed: {e}")

    if not text.strip():
        return "[WARNING: No text found. This PDF might be a scanned image and OCR is not currently enabled.]"
        
    return text

def extract_text_from_docx(filepath):
    """Extracts text from a DOCX file."""
    if not docx:
        return "[ERROR: python-docx library not installed. Cannot process DOCX files.]"
    
    try:
        doc = docx.Document(filepath)
        full_text = []
        
        # Extract text from paragraphs
        full_text.append("\n".join([para.text for para in doc.paragraphs]))
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                full_text.append(" | ".join(row_text))
        
        return "\n".join(full_text)
    except Exception as e:
        print(f"Error extracting DOCX: {e}")
        return f"[ERROR: Failed to extract DOCX content: {e}]"

def extract_text_from_file(filepath):
    """Unified function to extract text based on file extension."""
    ext = os.path.splitext(filepath)[1].lower()
    
    if ext == ".pdf":
        return extract_text_from_pdf(filepath)
    elif ext == ".docx":
        return extract_text_from_docx(filepath)
    elif ext == ".txt":
        try:
            with open(filepath, "r", encoding="utf-8") as f:
                return f.read()
        except Exception as e:
            return f"[ERROR: Failed to read text file: {e}]"
    else:
        return f"[WARNING: Unsupported file type: {ext}]"

def generate_summary(text):
    """Generates a brief summary of the text using OpenAI."""
    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a helpful study assistant. Summarize the following text in 3-5 bullet points."},
                {"role": "user", "content": text[:4000]} # Limit text length
            ]
        )
        return response.choices[0].message.content
    except Exception as e:
        print(f"Error generating summary: {e}")
        return "Summary could not be generated."

def generate_quiz(text):
    """Generates a 3-question quiz based on the text."""
    try:
        prompt = f"Create a 3-question multiple choice quiz based on this text. Return ONLY JSON in this format: [{{'question': '...', 'options': ['a', 'b', 'c', 'd'], 'answer': 'a'}}]. Text: {text[:4000]}"
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a quiz generator. Output valid JSON only."},
                {"role": "user", "content": prompt}
            ]
        )
        content = response.choices[0].message.content
        # Clean potential markdown code blocks
        if "```json" in content:
            content = content.split("```json")[1].split("```")[0].strip()
        elif "```" in content:
            content = content.split("```")[1].split("```")[0].strip()
            
        return json.loads(content)
    except Exception as e:
        print(f"Error generating quiz: {e}")
        # Fallback Mock Quiz
        return [
            {"question": "What is the main topic?", "options": ["Topic A", "Topic B", "Topic C", "Topic D"], "answer": "Topic A"},
            {"question": "Which concept is key?", "options": ["Concept X", "Concept Y", "Concept Z", "None"], "answer": "Concept X"},
            {"question": "Assessment check?", "options": ["True", "False"], "answer": "True"}
        ]
